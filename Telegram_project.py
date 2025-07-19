from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from docx.shared import Inches
from sentence_transformers import SentenceTransformer
import chromadb
import os
import threading
import tempfile
import subprocess
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import requests
import speech_recognition as sr
from PIL import Image
import torch
from transformers import CLIPProcessor, CLIPModel
import re

# === Configuration ===
BOT_TOKEN = "7969728700:AAEo5sqYN8RubzUfzkUUYo44SuK0LIgBRhE"
CHROMA_PATH = "/workspaces/Telegram-bot/chromaDB_data"
DOC_FILE = "/workspaces/Telegram-bot/Tech store table.docx"
IMAGE_DIR = os.path.join(os.path.dirname(DOC_FILE), "Products Images")
GROQ_API_KEY = "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"
GROQ_MODEL = "llama-3.3-70b-versatile"

# === ChromaDB & Model Setup ===
client = chromadb.Client()
collection = client.get_or_create_collection("tech_products")
image_collection = client.get_or_create_collection("image_products")
model = SentenceTransformer("all-MiniLM-L6-v2")

# === CLIP Model for Image Search ===
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

def get_image_embedding(img_path):
    try:
        image = Image.open(img_path).convert("RGB")
        image = image.resize((224, 224))
        inputs = clip_processor(images=image, return_tensors="pt")
        with torch.no_grad():
            embedding = clip_model.get_image_features(**inputs)
        return embedding[0].numpy()
    except Exception as e:
        print(f"❌ Error in image embedding: {e}")
        return None

def search_similar_image(img_path):
    try:
        query_embedding = get_image_embedding(img_path)
        if query_embedding is None:
            return None

        result = image_collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=1  # Only top match
        )

        if result and result["metadatas"] and len(result["metadatas"][0]) > 0:
            return result["metadatas"][0][0]  # Return the first metadata dict
        else:
            return None

    except Exception as e:
        print(f"❌ Image search error: {e}")
        return None



# === Smart DOCX Reader with Image Lookup from Folder ===
def extract_data_from_docx(file_path):
    doc = Document(file_path)
    contact_info = {"store_name": "Not found", "address": "Not found", "phone": "Not found", "website": "Not found"}

    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith("store name:"):
            contact_info["store_name"] = line.split(":", 1)[1].strip()
        elif low.startswith("location:"):
            contact_info["address"] = line.split(":", 1)[1].strip()
        elif low.startswith("phone:"):
            contact_info["phone"] = line.split(":", 1)[1].strip()
        elif low.startswith("web site:") or low.startswith("website:"):
            contact_info["website"] = line.split(":", 1)[1].strip()

    if not os.path.exists(IMAGE_DIR):
        os.makedirs(IMAGE_DIR)

    products = []
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        expected = ["product category", "product name", "product details", "product price", "stocks", "product id", "url"]
        if not all(col in headers for col in expected):
            continue

        idx = {col: headers.index(col) for col in expected}
        for row in table.rows[1:]:
            cells = row.cells
            product_id = cells[idx["product id"]].text.strip()  # Unique product ID
            image_path = None

            # Match image with the product_id
            for ext in [".jpg", ".jpeg", ".png"]:
                candidate = os.path.join(IMAGE_DIR, f"{product_id}{ext}")
                if os.path.exists(candidate):
                    image_path = candidate
                    break

            # Add product to the list
            products.append({
                "category": cells[idx["product category"]].text.strip(),
                "name":     cells[idx["product name"]].text.strip(),
                "details":  cells[idx["product details"]].text.strip(),
                "price":    cells[idx["product price"]].text.strip(),
                "stock":    cells[idx["stocks"]].text.strip(),
                "image":    image_path,  # Image path matching the product_id
                "url":      cells[idx["url"]].text.strip(),
            })

    return {"contact_info": contact_info, "products": products}

# === ChromaDB Functions ===
def clear_chroma_collection():
    try:
        all_ids = collection.get()["ids"]
        if all_ids:
            collection.delete(ids=all_ids)
            print("✅ Cleared old data from ChromaDB.")
    except Exception as e:
        print(f"❌ Failed to clear: {e}")

def upload_to_chroma(data):
    contact_info = data["contact_info"]
    contact_embed = model.encode([f"{contact_info['store_name']} | {contact_info['address']} | {contact_info['phone']} | {contact_info['website']}"])[0]

    # Add contact info to text-based collection
    collection.add(
        ids=["contact_info"],
        embeddings=[contact_embed.tolist()],
        documents=[f"{contact_info['store_name']} | {contact_info['address']} | {contact_info['phone']} | {contact_info['website']}"],
        metadatas=[contact_info]
    )

    for product in data["products"]:
        doc_text = f"{product['category']} - {product['name']} - {product['price']}"
        text_embedding = model.encode([doc_text])[0]

        try:
            # Add to text embedding collection
            collection.add(
                ids=[product["name"]],
                embeddings=[text_embedding.tolist()],
                documents=[doc_text],
                metadatas=[product]
            )
        except:
            pass

        # Add to image embedding collection if valid image exists
        if product.get("image") and os.path.exists(product["image"]):
            img_embedding = get_image_embedding(product["image"])
            if img_embedding is not None:
                try:
                    image_collection.add(
                        ids=[product["image"]],
                        embeddings=[img_embedding.tolist()],
                        documents=[product["name"]],
                        metadatas=[product]
                    )
                except Exception as e:
                    print(f"⚠️ Image add failed for {product['image']}: {e}")


# === Watchdog for .docx Changes ===
class DocxChangeHandler(FileSystemEventHandler):
    def on_modified(self, event):
        if event.src_path.lower().endswith(os.path.basename(DOC_FILE).lower()):
            print("📄 .docx changed, reloading...")
            try:
                clear_chroma_collection()
                data = extract_data_from_docx(DOC_FILE)
                upload_to_chroma(data)
                print("✅ Reload complete.")
            except Exception as e:
                print(f"❌ Reload failed: {e}")

def start_docx_watcher():
    observer = Observer()
    observer.schedule(DocxChangeHandler(), path=os.path.dirname(DOC_FILE), recursive=False)
    observer.start()
    threading.Thread(target=observer.join, daemon=True).start()

# === Helper Functions ===
def search_product(query):
    q = query.strip().lower()
    results = []
    for m in collection.get()["metadatas"]:
        if "name" in m and "category" in m:
            if q in m["name"].lower() or q in m["category"].lower():
                results.append(m)
    return results

def get_contact_info():
    try:
        meta = collection.get(ids=["contact_info"])["metadatas"]
        return meta[0] if meta else {}
    except:
        return {}

def ask_groq_with_context(user_query):
    try:
        all_data = collection.get()
        all_products = [m for m in all_data["metadatas"] if "name" in m]
        contact_info = [m for m in all_data["metadatas"] if "name" not in m]

        context = "Store Contact Info:\n"
        if contact_info:
            ci = contact_info[0]
            context += f"- Store Name: {ci.get('store_name', 'N/A')}\n"
            context += f"- Website: {ci.get('website', 'N/A')}\n"
            context += f"- Address: {ci.get('address', 'N/A')}\n"
            context += f"- Phone: {ci.get('phone', 'N/A')}\n"

        context += "\nAvailable Products:\n"
        for p in all_products:
            context += (
                f"- {p['name']} ({p['category']})\n"
                f"  • Details: {p['details']}\n"
                f"  • Price: {p['price']}\n"
                f"  • Stock: {p['stock']}\n"
                f"  • Image: {p['image']}\n"
                f"  • URL: {p['url']}\n\n"
            )

        full_prompt = (
            f"You are a helpful assistant for a tech store.\n"
            f"Use the following store info and products to answer the customer's question.\n\n"
            f"{context}\n"
            f"Customer: {user_query}"
        )

        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": GROQ_MODEL,
            "messages": [{"role": "user", "content": full_prompt}],
            "temperature": 0.7
        }

        res = requests.post(url, json=payload, headers=headers)
        res.raise_for_status()
        return res.json()["choices"][0]["message"]["content"].strip()

    except Exception as e:
        return f"❌ Error from LLM: {e}"

# === Telegram Handlers ===
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [['📍 Show Address', '📞 Show Phone'], ['🢍 Search by Product Name or Category']]
    markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text("👋 Welcome to Tech Store Bot!\nPlease choose an option:", reply_markup=markup)

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    ci = get_contact_info()

    if text == '📍 Show Address':
        await update.message.reply_text(f"🏠 {ci.get('address', 'Not found')}")
    elif text == '📞 Show Phone':
        await update.message.reply_text(f"📞 {ci.get('phone', 'Not found')}")
    elif text == '🢍 Search by Product Name or Category':
        await update.message.reply_text("🔍 Please enter the product name or category:")
        context.user_data["search_mode"] = True
    else:
        if context.user_data.get("search_mode", False):
            results = search_product(text)
            if results:
                for r in results:
                    caption = (
                        f"📦 Name: {r['name']}\n"
                        f"📂 Category: {r['category']}\n"
                        f"📝 Details:\n{r['details']}\n"
                        f"💵 Price: {r['price']}\n"
                        f"📊 Stock: {r['stock']}\n"
                        f"🔗 URL: {r['url']}"
                    )
                    if r["image"].startswith("http"):
                        await update.message.reply_photo(photo=r["image"], caption=caption)
                    elif os.path.exists(r["image"]):
                        await update.message.reply_photo(photo=open(r["image"], "rb"), caption=caption)
                    else:
                        await update.message.reply_text(caption)
            else:
                await update.message.reply_text("❌ Product not found. Asking assistant...")
                reply = ask_groq_with_context(text)

                image_matches = re.findall(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', reply, re.IGNORECASE)
                for img in image_matches:
                    if img.startswith("http"):
                        await update.message.reply_photo(photo=img)
                    elif os.path.exists(img):
                        await update.message.reply_photo(photo=open(img, "rb"))
                clean_reply = re.sub(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', '', reply).strip()
                if clean_reply:
                    await update.message.reply_text(clean_reply)

            context.user_data["search_mode"] = False
        else:
            reply = ask_groq_with_context(text)
            image_matches = re.findall(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', reply, re.IGNORECASE)
            for img in image_matches:
                if img.startswith("http"):
                    await update.message.reply_photo(photo=img)
                elif os.path.exists(img):
                    await update.message.reply_photo(photo=open(img, "rb"))
            clean_reply = re.sub(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', '', reply).strip()
            if clean_reply:
                await update.message.reply_text(clean_reply)


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        file = await update.message.voice.get_file()
        ogg_path = tempfile.mktemp(suffix=".ogg")
        wav_path = tempfile.mktemp(suffix=".wav")
        await file.download_to_drive(ogg_path)
        subprocess.run(['ffmpeg', '-i', ogg_path, wav_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)

        if text.strip():
            print(f"🗣️ Voice recognized: {text}")
            await update.message.reply_text(f"🎤 You said: {text}")
            reply = ask_groq_with_context(text)

            image_matches = re.findall(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', reply, re.IGNORECASE)
            for img in image_matches:
                if img.startswith("http"):
                    await update.message.reply_photo(photo=img)
                elif os.path.exists(img):
                    await update.message.reply_photo(photo=open(img, "rb"))
            clean_reply = re.sub(r'([A-Za-z]:[/\\][\w\s./\\-]+?\.(?:jpg|jpeg|png)|https?://[^\s]+?\.(?:jpg|jpeg|png))', '', reply).strip()
            if clean_reply:
                await update.message.reply_text(clean_reply)
        else:
            await update.message.reply_text("❌ Could not understand the voice.")
    except Exception as e:
        await update.message.reply_text(f"❌ Voice processing error: {e}")


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        # Get the image sent by user
        photo = update.message.photo[-1]
        image_file = await photo.get_file()
        img_path = tempfile.mktemp(suffix=".jpg")
        await image_file.download_to_drive(img_path)

        # Perform image similarity search using CLIP + Chroma
        match = search_similar_image(img_path)

        if match:
            product = match  # match is a dict with full product metadata

            caption = (
                f"📦 Name: {product['name']}\n"
                f"📂 Category: {product['category']}\n"
                f"📝 Details:\n{product['details']}\n"
                f"💵 Price: {product['price']}\n"
                f"📊 Stock: {product['stock']}\n"
                f"🔗 URL: {product['url']}"
            )

            if product["image"] and product["image"].startswith("http"):
                await update.message.reply_photo(photo=product["image"], caption=caption)
            elif product["image"] and os.path.exists(product["image"]):
                await update.message.reply_photo(photo=open(product["image"], "rb"), caption=caption)
            else:
                await update.message.reply_text(caption)
        else:
            await update.message.reply_text("❌ No similar product found.")

    except Exception as e:
        await update.message.reply_text(f"❌ Error in image processing: {e}")

#yo yo
# === Main Entrypoint ===
def main():
    print("🔄 Initial load of ChromaDB...")
    clear_chroma_collection()
    data = extract_data_from_docx(DOC_FILE)
    upload_to_chroma(data)

    print("👀 Watching for .docx changes...")
    start_docx_watcher()

    print("🤖 Starting Telegram bot...")
    app = Application.builder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.run_polling()

if __name__ == "__main__":
    main()
